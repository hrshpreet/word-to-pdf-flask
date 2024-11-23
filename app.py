from flask import Flask, render_template, request, jsonify
import subprocess
import os
from docx import Document

# subprocess.run(['soffice', '--headless', '--convert-to', 'pdf', 'email.docx', '--outdir', 'outputs'], check=True)

app = Flask(__name__)

app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024


UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = os.path.join(os.getcwd(), 'static', 'converted_pdfs')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


def extract_metadata(file_path):
    doc = Document(file_path)
    
    # Basic metadata
    metadata = {
        "author": doc.core_properties.author,
        "title": doc.core_properties.title,
        "subject": doc.core_properties.subject,
        "keywords": doc.core_properties.keywords,
        "created": doc.core_properties.created,
        "modified": doc.core_properties.modified,
        "word_count": sum(len(paragraph.text.split()) for paragraph in doc.paragraphs),
        "paragraph_count": len(doc.paragraphs),
        "size": os.path.getsize(file_path)
    }
    return metadata



@app.route('/')
def index():
    return render_template('index.html')



@app.route('/convert', methods=['POST'])
def convert_to_pdf():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400

    file = request.files['file']
    
    if not file.filename.lower().endswith('.docx'):
        return jsonify({'error': 'Invalid file type. Only .docx files are allowed.'}), 400

    input_path = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(input_path)
    metadata = extract_metadata(input_path)
    
    output_filename = file.filename.rsplit('.', 1)[0] + '.pdf'
    output_path = os.path.join(OUTPUT_FOLDER, output_filename)

    try:
        subprocess.run(['soffice', '--headless', '--convert-to', 'pdf', input_path, '--outdir', OUTPUT_FOLDER], check=True)
        return jsonify({
            'output': f'/static/converted_pdfs/{output_filename}',
            'metadata': metadata
        }), 200
        
    except subprocess.CalledProcessError as e:
        return jsonify({'error': str(e)}), 500
    
    finally:
        os.remove(input_path)
        

@app.errorhandler(413)
def request_entity_too_large(error):
    return jsonify({'error': 'File is too large. Maximum size allowed is 10 MB.'}), 413



if __name__ == '__main__':
    app.run(debug=True)



